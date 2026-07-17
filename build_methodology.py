from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import docx.oxml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.color.rgb = RGBColor(0x0E, 0x0E, 0x0E)

sections = doc.sections
for section in sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(10)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Arial'
        shading = docx.oxml.parse_xml(f'<w:shd {qn("w:fill")}="0E0E0E" {qn("w:val")}="clear" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Arial'
    doc.add_paragraph()

# ═══ TITLE ═══
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("YouTube Insights Report\n")
run.bold = True
run.font.size = Pt(26)
run = p.add_run("Full Methodology & Audit Documentation\n")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x8A, 0x84, 0x7A)
run = p.add_run("Virgin Music Group — June 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x8A, 0x84, 0x7A)

doc.add_paragraph()
add_warning("CRITICAL FINDING: This report has no analytical pipeline. Every number is a hardcoded constant. The report cannot be independently reproduced from the codebase alone.")

doc.add_page_break()

# ═══ 1. ARCHITECTURE ═══
doc.add_heading('1. System Architecture', level=1)
doc.add_paragraph('The YouTube Insights report is a Next.js 14.2.5 static website deployed to Vercel. It contains 4 source files and no API routes, no database connections, and no external data fetching.')

doc.add_heading('1.1 File Inventory', level=2)
add_table(
    ["File", "Lines", "Role"],
    [
        ["src/data/insights.ts", "491", "All data. Every number is a hardcoded constant."],
        ["src/app/page.tsx", "1,869", "All React components. Contains duplicate data in VisualProof components."],
        ["src/app/layout.tsx", "23", "HTML shell, Inter font, page metadata."],
        ["src/app/globals.css", "89", "CSS variables, nav styles, card animations."],
    ]
)

doc.add_heading('1.2 Data Flow', level=2)
doc.add_paragraph('The data flow is a two-hop pipeline with zero transformation: Static constants in insights.ts are imported by page.tsx and rendered directly. There is no computation, aggregation, filtering, or sorting.')
add_warning('There are zero API calls. No YouTube Data API. No Watcher API. No Supabase. The comment at insights.ts line 4 says "Future: connect to Watcher API / Supabase for live data" — this has not been implemented.')

doc.add_page_break()

# ═══ 2. DATA SOURCES ═══
doc.add_heading('2. Data Sources', level=1)
doc.add_heading('2.1 What The Report Claims', level=2)
doc.add_paragraph('The report presents itself as an analysis of 138 YouTube channels across 3,554 videos, with 82 Virgin-managed and 56 market benchmark channels, and 7 manually reviewed campaigns.')

doc.add_heading('2.2 What Exists In The Codebase', level=2)
doc.add_paragraph('None of these datasets exist. The numbers 138, 3554, 82, 56, and 7 are hardcoded string literals in the headlineStats array (insights.ts lines 8-14).')
doc.add_paragraph('The only data that exists: 21 video assets across 6 case studies, 6 campaign groupings (manually curated), 6 channel-level records (limited fields), 6 benchmark comparison rows, 4 follow-up segment percentages, and 4 diversity statistics.')

doc.add_heading('2.3 Where The Numbers Came From', level=2)
doc.add_paragraph('The numbers were generated in prior Cowork sessions using external tools. The task history (171+ completed tasks) confirms YouTube Data API calls, Watcher API queries, Python backfill scripts for 142 channels, and multiple report generations. The final numbers were manually entered into insights.ts.')
add_warning('The external analysis pipeline is not part of this codebase. The raw data, scripts, API credentials, and intermediate computations do not exist here.')

doc.add_page_break()

# ═══ 3. CLASSIFICATION RULES ═══
doc.add_heading('3. Classification Rules', level=1)

doc.add_heading('3.1 Channel Health Classifications', level=2)
doc.add_paragraph('The report uses channel health classifications: Growing, Weak Conversion, Underfed, Cold. These are NOT computed in this codebase. They are hardcoded strings assigned per case study.')
doc.add_paragraph('A WatcherBadge component (page.tsx lines 505-531) maps strings to colors: GROWING to green, WEAK_CONVERSION to yellow, UNDERFED to red, COLD to grey. No classification algorithm, scoring function, or threshold definitions exist.')
add_warning('"Healthy" and "Dormant" classifications do not exist anywhere in this codebase.')

doc.add_heading('3.2 Format Classification', level=2)
doc.add_paragraph('Video formats are manually assigned via the content_type field. No format detection algorithm exists — no title parsing, no duration check, no YouTube metadata analysis.')
add_table(
    ["content_type", "Display Name", "Used By"],
    [
        ["official_video", "Official Video / MV", "All 6 artists"],
        ["acoustic", "Acoustic", "mary in the junkyard"],
        ["visualiser", "Visualiser", "mary in the junkyard"],
        ["bts", "Behind the Scenes", "GENER8ION, K-Trap"],
        ["performance", "Performance", "GENER8ION"],
        ["trailer", "Trailer", "GENER8ION"],
        ["live", "Live", "The Snuts"],
        ["lyric_video", "Lyric Video", "The Snuts (x2)"],
    ]
)

doc.add_heading('3.3 Short Detection', level=2)
doc.add_paragraph('Shorts are referenced in two places (followUpSegments: 35%, shortsLayer description) but NO individual Short appears in any asset list. There is no duration threshold, no title pattern match, no YouTube metadata check. The 35% figure was computed externally.')

doc.add_page_break()

# ═══ 4. CAMPAIGN RULES ═══
doc.add_heading('4. Campaign Detection & Grouping Rules', level=1)

doc.add_heading('4.1 Campaign Detection', level=2)
doc.add_paragraph('Campaigns are not algorithmically detected. Each campaign is a manually curated array of assets in insights.ts. There is no proximity-based grouping, date windowing, or algorithmic detection.')

doc.add_heading('4.2 Campaign Boundaries', level=2)
doc.add_paragraph('Campaign Start: Implicit — the asset with days_from_hero: 0. No explicit start date stored. The release_date field refers to the album release, which may differ.')
doc.add_paragraph('Campaign End: Implicit — the highest days_from_hero value. No explicit end date or maximum gap rule.')

doc.add_heading('4.3 Asset Inclusion/Exclusion', level=2)
doc.add_paragraph('Undocumented. Each case study has 3-5 manually selected assets. No documentation of which assets were considered but excluded, or the inclusion criteria.')

doc.add_heading('4.4 Campaign Moments vs Listed Assets', level=2)
add_warning('K-Trap claims 17 campaign moments but lists 4 assets. The Snuts claims 22 campaign moments but lists 4 assets. The source of the higher counts is not documented. These numbers likely come from the full channel analysis (all uploads, not just selected campaign assets) but the full upload list is not in the codebase.')

doc.add_page_break()

# ═══ 5. BENCHMARK METHODOLOGY ═══
doc.add_heading('5. Benchmark Methodology', level=1)

doc.add_heading('5.1 Percentile Rankings', level=2)
doc.add_paragraph('Rankings use: Rank / Total Channels. Example: Rank 13 of 138 = 13/138 = 9.4%, displayed as "Top 9%".')
add_warning('Different artists are ranked on DIFFERENT metrics: Tove Lo on "multi-format campaign", K-Trap on "format diversity", The Snuts on "follow-up support", GENER8ION on "efficiency". These are not comparable rankings from a single leaderboard.')

doc.add_heading('5.2 Virgin vs Market Comparison', level=2)
add_table(
    ["Metric", "Virgin", "Market", "Gap", "Actionable"],
    [
        ["Live Sessions", "48%", "63%", "-15 pts", "Yes"],
        ["Lyric Videos", "44%", "54%", "-10 pts", "Yes"],
        ["Multi-Format Campaign", "13%", "18%", "-5 pts", "Yes"],
        ["Format Variety", "4.8 avg", "4.6 avg", "+0.2", "No"],
        ["Visualisers", "65%", "61%", "+4 pts", "No"],
        ["BTS Content", "26%", "30%", "-4 pts", "No"],
    ]
)

doc.add_heading('5.3 Consistency Check: 85% vs 76%', level=2)
doc.add_paragraph('toveLoEvidence.benchmarkContext states: "85% of campaigns go silent after Day 7." followUpSegments shows 41% silent + 35% shorts only = 76%. The 85% figure likely uses a different definition ("no longform follow-up after Day 7" vs "completely silent") but the difference is not explained in the codebase.')

doc.add_page_break()

# ═══ 6. INCONSISTENCIES ═══
doc.add_heading('6. Identified Inconsistencies', level=1)
add_table(
    ["#", "Inconsistency", "Detail"],
    [
        ["1", "GENER8ION total views", "Claimed 23.6M. Visible assets sum to 15.96M. Gap of 7.64M."],
        ["2", "K-Trap format types", "Claimed 8. Only 2 visible (official_video, bts)."],
        ["3", "K-Trap campaign moments", "Claimed 17. Only 4 assets listed."],
        ["4", "Snuts campaign moments", "Claimed 22. Only 4 assets listed."],
        ["5", "85% vs 76% silent", "Different definitions, gap unexplained."],
        ["6", "7 campaigns vs 6 visible", "7th campaign not identified."],
        ["7", "Data duplication", "Case study data in insights.ts AND in VisualProof components."],
    ]
)

doc.add_page_break()

# ═══ 7. REPRODUCIBILITY ═══
doc.add_heading('7. Reproducibility Requirements', level=1)
doc.add_paragraph('For another analyst or LLM to independently reproduce this report, the following exports and documentation are needed:')

items = [
    ("Data Exports", [
        "Full 138-channel dataset as CSV with all fields used for classification and ranking",
        "Full 3,554-video dataset as CSV with format types, campaign assignments, and view counts at capture date",
        "Classification algorithm documentation with thresholds and scoring functions",
        "Full ranked list for each ranking criterion used (multi-format, format diversity, follow-up support, efficiency)",
    ]),
    ("Methodology Documentation", [
        "Follow-up analysis methodology: definitions of silent, shorts-only, limited longform, multi-format",
        "Format detection rules: how Shorts are identified, how each format type is assigned",
        "Campaign grouping criteria: maximum gap, inclusion/exclusion rules",
        "View count capture methodology: when recorded, lifetime vs period",
    ]),
    ("Code Changes", [
        "Store video IDs and YouTube URLs for all case study assets (currently all null)",
        "Resolve data duplication between insights.ts and VisualProof components",
        "Add actual dates alongside day offsets",
        "Connect to live APIs to replace hardcoded constants",
    ]),
]

for section_title, bullets in items:
    doc.add_heading(section_title, level=2)
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ═══ 8. CONCLUSION ═══
doc.add_heading('8. Conclusion', level=1)
doc.add_paragraph('This audit reveals that the YouTube Insights report is a static presentation, not an analytical pipeline. Every number shown to the reader was manually entered as a constant. The underlying analysis was performed in prior sessions using external tools and APIs, but that work product does not exist in this codebase.')
doc.add_paragraph('The report makes claims about 138 channels and 3,554 videos, but contains data for only 6 artists and 21 assets. The gap between what is claimed and what is verifiable is significant.')
doc.add_paragraph('Several internal inconsistencies were identified. These may have valid explanations from the external analysis, but those explanations are not documented here.')
doc.add_paragraph('The case study timelines were validated against YouTube in a prior session (the Critical Validation Pass) and are likely accurate. The editorial observations and learnings are subjective interpretations that cannot be independently verified.')
p = doc.add_paragraph('To make this report fully auditable, the external datasets and methodology need to be exported and stored alongside the presentation code.')
p.runs[0].bold = True

output_path = "/sessions/amazing-jolly-cannon/mnt/Desktop/YouTube_Insights_Methodology.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
